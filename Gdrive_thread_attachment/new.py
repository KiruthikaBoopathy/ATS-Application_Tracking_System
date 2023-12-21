import concurrent.futures
import io
import os
import tempfile
import time
from email.header import decode_header
from email.utils import parsedate_to_datetime
import email
import imaplib
import html2text
import json
import re
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import fitz  # PyMuPDF
import openpyxl
from docx import Document


class EmailExtractor:
    def __init__(self, imap_user_email, imap_password):
        self.imap_user_email = imap_user_email
        self.imap_password = imap_password

    def get_emails(self):
        imap_server = "imap.gmail.com"
        imap_port = 993

        imap_conn = imaplib.IMAP4_SSL(imap_server, imap_port)
        imap_conn.login(self.imap_user_email, self.imap_password)
        imap_conn.select('inbox')

        _, message_ids = imap_conn.search(None, "UNSEEN")

        emails = []

        for message_id in message_ids[0].split():
            try:
                _, data = imap_conn.fetch(message_id, "(RFC822)")
                raw_email = data[0][1]
                email_message = email.message_from_bytes(raw_email)

                subject = email_message["Subject"]
                sender = email.utils.parseaddr(email_message["From"])[0]
                body1 = ""
                attachments = []

                date_received = parsedate_to_datetime(email_message["Date"])

                if email_message.is_multipart():
                    for part in email_message.walk():
                        if part.get_content_type() == "text/plain":
                            body1 = part.get_payload(decode=True).decode()
                        elif part.get_content_type() == "text/html":
                            body1 = html2text.html2text(part.get_payload(decode=True).decode())
                        elif part.get_content_maintype() != 'multipart' and part.get(
                                'Content-Disposition') is not None:
                            filename = part.get_filename()
                            if filename:
                                filename, encoding = decode_header(filename)[0]
                                if isinstance(filename, bytes):
                                    filename = filename.decode(encoding or "utf-8")
                                attachment_data = part.get_payload(decode=True)
                                attachments.append({
                                    "filename": filename,
                                    "data": attachment_data
                                })

                email_data = {
                    "subject": subject,
                    "sender": sender,
                    "body": body1,
                    "date_received": date_received,
                    "email_id": message_id,
                    "attachments": attachments
                }

                emails.append(email_data)
            except Exception as e:
                print(f"Error processing email: {str(e)}")

                imap_conn.store(message_id, '-FLAGS', '(\Seen)')

        imap_conn.close()

        return emails

    def process_emails(self, emails):
        with concurrent.futures.ThreadPoolExecutor() as executor:
            executor.map(self.op_process_mail, emails)

    def op_process_mail(self, email_data):
        global attachment_info, details
        print('\n')
        print('Source:Gmail')
        print("Subject:", email_data["subject"])
        print("Sender:", email_data["sender"])
        print("Date Received:", email_data["date_received"])

        if email_data["attachments"]:
            for attachment in email_data["attachments"]:
                filename = attachment["filename"]
                content = attachment["data"]

                file_format = filename.split('.')[-1].lower()

                if file_format == 'xlsx':
                    wb = openpyxl.load_workbook(io.BytesIO(attachment["data"]), read_only=True)
                    pdf_content = ""
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        for row in sheet.iter_rows():
                            row_values = [str(cell.value) for cell in row if cell.value is not None]
                            pdf_content += " ".join(row_values) + "\n"
                        pdf_content += "\n"

                    attachment_info = {
                        "filename": filename,
                        "content_type": "text_from_xlsx",
                        "content": pdf_content
                    }

                elif file_format == 'pdf':
                    pdf_document = fitz.open(stream=io.BytesIO(attachment["data"]), filetype="pdf")

                    pdf_content = ""

                    for page_number in range(pdf_document.page_count):
                        page = pdf_document[page_number]
                        pdf_content += page.get_text()


                    attachment_info = {
                        "filename": filename,
                        "content_type": "text_from_pdf",
                        "content": pdf_content
                    }

                elif file_format == 'docx':
                    doc = Document(io.BytesIO(attachment["data"]))
                    pdf_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

                    attachment_info = {
                        "filename": filename,
                        "content_type": "text_from_docx",
                        "content": pdf_content
                    }

                elif file_format == 'py':
                    pdf_content = attachment["data"].decode('utf-8')
                    attachment_info = {
                        "filename": filename,
                        "content_type": "text_from_python",
                        "content": pdf_content
                    }

            if attachment_info:
                print("Attachments:")
                print(json.dumps(attachment_info, indent=2))
            else:
                print("No Attachments")

        else:
            if email_data["body"]:
                print("Body:", email_data["body"])
            else:
                print("No Body content")


class DriveExtractor:
    def __init__(self, root_directory_id, credentials_file):
        self.root_directory_id = root_directory_id
        self.credentials_file = credentials_file

    def authenticate_drive_service(self):
        SCOPES = ['https://www.googleapis.com/auth/drive.metadata.readonly']

        credentials = None
        if os.path.exists('token.json'):
            credentials = Credentials.from_authorized_user_file('token.json')
        if not credentials or not credentials.valid:
            if credentials and credentials.expired and credentials.refresh_token:
                credentials.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(self.credentials_file, SCOPES)
                credentials = flow.run_local_server(port=0)

            with open('token.json', 'w') as token:
                token.write(credentials.to_json())

        drive_service = build('drive', 'v3', credentials=credentials)
        return drive_service

    def extract_content_from_drive_folder(self, folder_id, drive_service):
        results = drive_service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            pageSize=10, fields="nextPageToken, files(id, name, mimeType)").execute()
        files = results.get('files', [])

        if not files:
            print(f'No files found in the folder with ID: {folder_id}')
        else:

            for file in files:
                file_id = file['id']
                file_name = file['name']
                mime_type = file['mimeType']

                request = drive_service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False

                while not done:
                    status, done = downloader.next_chunk()

                if mime_type == 'application/pdf':
                    pdf_content = self.extract_pdf_content(fh)
                    drive_file_data = {
                                "Source": 'Google drive',
                                "file_name": file_name,
                                "file_type": "pdf",
                                "content_type": "text_from_pdf",
                                "content": pdf_content
                    }
                    print( f'Final Output:\n{drive_file_data}\n{"=" * 50}\n' )

                elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    docx_content = self.extract_docx_content(fh)
                    drive_file_data = {
                                "Source": 'Google drive',
                                "file_name": file_name,
                                "file_type": "docx",
                                "content_type": "text_from_docx",
                                "content": docx_content
                    }
                    print( f'Final Output:\n{drive_file_data}\n{"=" * 50}\n' )

                elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                    excel_content = self.extract_excel_content(fh)
                    drive_file_data = {

                        "Source":'Google drive',
                        "file_name": file_name,
                        "file_type": "xlsx",
                        "content_type": "text_from_xlsx",
                        "content": excel_content
                    }
                    print( f'Final Output:\n{drive_file_data}\n{"=" * 50}\n' )

                else:
                    print(f'Skipping unsupported file type: {mime_type}')
                    continue

    @staticmethod
    def extract_pdf_content(file_content):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(file_content.getvalue())
            temp_pdf_path = temp_pdf.name

            pdf_document = fitz.open(temp_pdf_path)
            pdf_text = ""

            for page_number in range(pdf_document.page_count):
                page = pdf_document[page_number]
                pdf_text += page.get_text()

            return pdf_text

    @staticmethod
    def extract_excel_content(file_content):
        workbook = openpyxl.load_workbook(file_content)
        excel_text = ""

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            excel_text += f"Sheet: {sheet_name}\n"

            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) for cell in row)
                excel_text += f"{row_text}\n"

            excel_text += "\n"

        return excel_text

    @staticmethod
    def extract_docx_content(file_content):
        doc = Document(file_content)
        docx_text = ""

        for paragraph in doc.paragraphs:
            docx_text += paragraph.text + "\n"

        return docx_text


if __name__ == "__main__":
    imap_user_email = "kiruthika.b@vrdella.com"
    imap_password = "renm kixf nlxy avbx"
    email_extractor = EmailExtractor(imap_user_email, imap_password)

    drive_credentials_file = r'C:\Users\Vrdella\Downloads\gdrive_credentials.json'
    root_directory_id = 'root'
    drive_extractor = DriveExtractor(root_directory_id, drive_credentials_file)

    # Extract emails
    start_email_time = time.time()
    emails = email_extractor.get_emails()
    email_extractor.process_emails(emails)
    end_email_time = time.time()

    print(f"Email extraction time: {end_email_time - start_email_time} seconds")

    # Extract Google Drive content
    start_drive_time = time.time()
    drive_service = drive_extractor.authenticate_drive_service()

    results = drive_service.files().list(
        q=f"'{root_directory_id}' in parents and trashed=false and mimeType='application/vnd.google-apps.folder'",
        pageSize=10, fields="nextPageToken, files(id, name)").execute()
    folders = results.get('files', [])

    if not folders:
        print(f'No folders found in the root directory')
    else:

        for folder in folders:
            folder_id = folder['id']
            folder_name = folder['name']

            # Process each folder in a separate thread
            with concurrent.futures.ThreadPoolExecutor() as executor:
                futures = [
                    executor.submit(drive_extractor.extract_content_from_drive_folder, folder_id, drive_service)]

                # Wait for all tasks to complete
                concurrent.futures.wait(futures)

    end_drive_time = time.time()
    print(f"Google Drive extraction time: {end_drive_time - start_drive_time} seconds")
